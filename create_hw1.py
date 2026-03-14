from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# --- Page margins ---
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(3)
section.right_margin = Cm(1.5)

# --- Styles helpers ---
def set_font(run, size=12, bold=False, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_font(run, 14, bold=True)
    elif level == 2:
        set_font(run, 13, bold=True)
    else:
        set_font(run, 12, bold=True)
    return p

def add_paragraph(text='', bold=False, indent=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    set_font(run, 12, bold=bold)
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, 12)
    return p

def add_numbered(text, num):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(f'{num}. {text}')
    set_font(run, 12)
    return p

def shade_cell(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_table(headers, rows, key_cols=None, caption=None):
    """headers: list of str, rows: list of lists, key_cols: set of col indices that are PK"""
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        shade_cell(cell, 'D9D9D9')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        is_key = key_cols and i in key_cols
        set_font(run, 11, bold=True, color=(255, 0, 0) if is_key else None)

    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            is_key = key_cols and ci in key_cols
            set_font(run, 11, bold=True if is_key else False,
                     color=(255, 0, 0) if is_key else None)

    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(caption)
        set_font(run, 11, bold=False)
        run.font.italic = True

    return table

# ============================================================
# TITLE
# ============================================================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('Домашнее задание №1')
set_font(title_run, 16, bold=True)
doc.add_paragraph()

# ============================================================
# ВВЕДЕНИЕ
# ============================================================
add_heading('Введение', level=1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Cm(1.25)
run = p.add_run(
    'Данная домашняя работа посвящена проектированию базы данных для информационно-аналитического '
    'приложения. '
)
set_font(run, 12)
run2 = p.add_run('Предметная область: ')
set_font(run2, 12, bold=True)
run3 = p.add_run('рабочее место менеджера по подбору персонала (учёт кандидатов, ведение базы вакансий, '
                 'анализ результатов собеседований). ')
set_font(run3, 12)
run4 = p.add_run('Потенциальные пользователи: ')
set_font(run4, 12, bold=True)
run5 = p.add_run('HR-менеджер, рекрутер, руководитель HR-отдела.')
set_font(run5, 12)

# ============================================================
# СОЗДАНИЕ ИНФОРМАЦИОННОЙ ОСНОВЫ
# ============================================================
doc.add_paragraph()
add_heading('Создание информационной основы для приложений', level=1)

add_heading('Первый этап: Модель предметной области', level=2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Cm(1.25)
run = p.add_run('В ходе анализа предметной области были выделены следующие закономерности:')
set_font(run, 12)

bullets = [
    'Кандидаты откликаются на открытые вакансии компании.',
    'В компании существует несколько отделов, в каждый из которых может требоваться новый сотрудник.',
    'Каждая вакансия закреплена за одним конкретным отделом.',
    'На одну вакансию может претендовать несколько кандидатов.',
    'Один кандидат может проходить собеседования на разные вакансии.',
    'Каждое собеседование проводится ровно одним HR-менеджером, имеет конкретную дату и результат.',
    'Запись о собеседовании однозначно определяется парой: кандидат и вакансия, на которую он претендует.',
]
for b in bullets:
    add_bullet(b)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p2.paragraph_format.first_line_indent = Cm(1.25)
r = p2.add_run('В ходе дополнительного уточнения того, какие данные необходимо учитывать, выяснилось следующее:')
set_font(r, 12)

details = [
    'О каждом кандидате необходимо хранить уникальный номер (Н_КАНД), ФИО и номер телефона.',
    'О каждом отделе хранится уникальный номер (Н_ОТД) и название.',
    'Каждая вакансия имеет уникальный номер (Н_ВАК) и наименование.',
    'О HR-менеджере хранится уникальный номер (Н_МЕН) и ФИО.',
    'Дата и результат относятся непосредственно к записи о собеседовании.',
]
for i, d in enumerate(details, 1):
    add_numbered(d, i)

# --- Схема начального отношения ---
doc.add_paragraph()
add_heading('Схема начального отношения', level=2)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p3.paragraph_format.first_line_indent = Cm(1.25)
r = p3.add_run(
    'На основании модели предметной области составим начальное отношение '
    'КАНДИДАТЫ_ВАКАНСИИ_ОТДЕЛЫ_МЕНЕДЖЕРЫ_СОБЕСЕДОВАНИЯ со следующими атрибутами:'
)
set_font(r, 12)

attrs = [
    ('Н_КАНД', 'номер кандидата в базе'),
    ('ФИО_КАНД', 'фамилия, имя и отчество кандидата'),
    ('ТЕЛ_КАНД', 'номер телефона кандидата'),
    ('Н_ВАК', 'номер открытой вакансии'),
    ('НАЗВ_ВАК', 'наименование вакансии'),
    ('Н_ОТД', 'номер отдела, в который ищут сотрудника'),
    ('НАЗВ_ОТД', 'название отдела'),
    ('Н_МЕН', 'номер HR-менеджера'),
    ('ФИО_МЕН', 'ФИО HR-менеджера'),
    ('ДАТА', 'дата проведения собеседования'),
    ('РЕЗУЛЬТАТ', 'итог собеседования'),
]
for abbr, desc in attrs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.25)
    r1 = p.add_run(f'{abbr}')
    set_font(r1, 12, bold=True)
    r2 = p.add_run(f' — {desc}')
    set_font(r2, 12)

p_note = doc.add_paragraph()
p_note.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p_note.paragraph_format.first_line_indent = Cm(1.25)
r = p_note.add_run('Красным цветом выделены атрибуты первичного ключа.')
set_font(r, 12, color=(255, 0, 0))

doc.add_paragraph()
add_table(
    headers=['Н_КАНД', 'ФИО_КАНД', 'ТЕЛ_КАНД', 'Н_ВАК', 'НАЗВ_ВАК', 'Н_ОТД', 'НАЗВ_ОТД', 'Н_МЕН', 'ФИО_МЕН', 'ДАТА', 'РЕЗУЛЬТАТ'],
    rows=[
        [1, 'Иванов И.И.', '111-22-33', 1, 'Разработчик', 1, 'IT-отдел', 1, 'Сидоров А.А.', '10.10.23', 'Оффер'],
        [1, 'Иванов И.И.', '111-22-33', 2, 'Аналитик', 2, 'Аналитика', 2, 'Петрова В.В.', '12.10.23', 'Отказ'],
        [2, 'Смирнов П.П.', '222-33-44', 1, 'Разработчик', 1, 'IT-отдел', 1, 'Сидоров А.А.', '13.10.23', 'В резерв'],
    ],
    key_cols={0, 3},
    caption='Таблица 1. Исходное отношение КАНДИДАТЫ_ВАКАНСИИ_ОТДЕЛЫ_МЕНЕДЖЕРЫ_СОБЕСЕДОВАНИЯ'
)

# ============================================================
# НОРМАЛИЗАЦИЯ
# ============================================================
doc.add_paragraph()
add_heading('Второй этап: Нормализация', level=1)

# --- 1НФ ---
add_heading('Первая нормальная форма (1НФ)', level=2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Cm(1.25)
r = p.add_run(
    'Отношение уже находится в 1НФ: все кортежи уникальны, атрибуты атомарны, порядок строк и столбцов не '
    'имеет значения. В качестве первичного ключа берём пару атрибутов '
)
set_font(r, 12)
r2 = p.add_run('{Н_КАНД, Н_ВАК}')
set_font(r2, 12, bold=True, color=(255, 0, 0))
r3 = p.add_run('.')
set_font(r3, 12)

# --- 2НФ ---
doc.add_paragraph()
add_heading('Вторая нормальная форма (2НФ)', level=2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Cm(1.25)
r = p.add_run(
    'Отношение НЕ находится во 2НФ: существуют не ключевые атрибуты, зависящие лишь от части составного ключа. '
    'Зависимости:'
)
set_font(r, 12)

deps = [
    'Н_КАНД → ФИО_КАНД, ТЕЛ_КАНД',
    'Н_ВАК → НАЗВ_ВАК, Н_ОТД, НАЗВ_ОТД',
]
for d in deps:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(2)
    r = p.add_run(d)
    set_font(r, 12, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Cm(1.25)
r = p.add_run('Декомпозируем исходное отношение на три:')
set_font(r, 12)

doc.add_paragraph()
p = doc.add_paragraph()
r1 = p.add_run('1. КАНДИДАТЫ')
set_font(r1, 12, bold=True)
r2 = p.add_run('  (Первичный ключ: ')
set_font(r2, 12)
r3 = p.add_run('Н_КАНД')
set_font(r3, 12, bold=True, color=(255, 0, 0))
r4 = p.add_run(')')
set_font(r4, 12)
add_table(
    headers=['Н_КАНД', 'ФИО_КАНД', 'ТЕЛ_КАНД'],
    rows=[
        [1, 'Иванов И.И.', '111-22-33'],
        [2, 'Смирнов П.П.', '222-33-44'],
    ],
    key_cols={0},
    caption='Таблица 2. Отношение КАНДИДАТЫ'
)

doc.add_paragraph()
p = doc.add_paragraph()
r1 = p.add_run('2. ВАКАНСИИ_ОТДЕЛЫ')
set_font(r1, 12, bold=True)
r2 = p.add_run('  (Первичный ключ: ')
set_font(r2, 12)
r3 = p.add_run('Н_ВАК')
set_font(r3, 12, bold=True, color=(255, 0, 0))
r4 = p.add_run(')')
set_font(r4, 12)
add_table(
    headers=['Н_ВАК', 'НАЗВ_ВАК', 'Н_ОТД', 'НАЗВ_ОТД'],
    rows=[
        [1, 'Разработчик', 1, 'IT-отдел'],
        [2, 'Аналитик', 2, 'Аналитика'],
    ],
    key_cols={0},
    caption='Таблица 3. Отношение ВАКАНСИИ_ОТДЕЛЫ'
)

doc.add_paragraph()
p = doc.add_paragraph()
r1 = p.add_run('3. СОБЕСЕДОВАНИЯ_МЕНЕДЖЕРЫ')
set_font(r1, 12, bold=True)
r2 = p.add_run('  (Первичный ключ: ')
set_font(r2, 12)
r3 = p.add_run('{Н_КАНД, Н_ВАК}')
set_font(r3, 12, bold=True, color=(255, 0, 0))
r4 = p.add_run(')')
set_font(r4, 12)
add_table(
    headers=['Н_КАНД', 'Н_ВАК', 'Н_МЕН', 'ФИО_МЕН', 'ДАТА', 'РЕЗУЛЬТАТ'],
    rows=[
        [1, 1, 1, 'Сидоров А.А.', '10.10.23', 'Оффер'],
        [1, 2, 2, 'Петрова В.В.', '12.10.23', 'Отказ'],
        [2, 1, 1, 'Сидоров А.А.', '13.10.23', 'В резерв'],
    ],
    key_cols={0, 1},
    caption='Таблица 4. Отношение СОБЕСЕДОВАНИЯ_МЕНЕДЖЕРЫ'
)

# --- 3НФ ---
doc.add_paragraph()
add_heading('Третья нормальная форма (3НФ)', level=2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Cm(1.25)
r = p.add_run(
    'Отношения ВАКАНСИИ_ОТДЕЛЫ и СОБЕСЕДОВАНИЯ_МЕНЕДЖЕРЫ не находятся в 3НФ из-за транзитивных зависимостей не ключевых атрибутов:'
)
set_font(r, 12)

trans = [
    'Н_ВАК → Н_ОТД → НАЗВ_ОТД',
    '{Н_КАНД, Н_ВАК} → Н_МЕН → ФИО_МЕН',
]
for t in trans:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(2)
    r = p.add_run(t)
    set_font(r, 12, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Cm(1.25)
r = p.add_run('Декомпозируем их дальше, получая итоговые 5 справочников:')
set_font(r, 12)

doc.add_paragraph()
p = doc.add_paragraph()
r1 = p.add_run('1. КАНДИДАТЫ')
set_font(r1, 12, bold=True)
r2 = p.add_run('  (Первичный ключ: ')
set_font(r2, 12)
r3 = p.add_run('Н_КАНД')
set_font(r3, 12, bold=True, color=(255, 0, 0))
r4 = p.add_run(') — без изменений из 2НФ')
set_font(r4, 12)
add_table(
    headers=['Н_КАНД', 'ФИО_КАНД', 'ТЕЛ_КАНД'],
    rows=[
        [1, 'Иванов И.И.', '111-22-33'],
        [2, 'Смирнов П.П.', '222-33-44'],
    ],
    key_cols={0},
    caption='Таблица 5. Отношение КАНДИДАТЫ'
)

doc.add_paragraph()
p = doc.add_paragraph()
r1 = p.add_run('2. ОТДЕЛЫ')
set_font(r1, 12, bold=True)
r2 = p.add_run('  (Первичный ключ: ')
set_font(r2, 12)
r3 = p.add_run('Н_ОТД')
set_font(r3, 12, bold=True, color=(255, 0, 0))
r4 = p.add_run(')')
set_font(r4, 12)
add_table(
    headers=['Н_ОТД', 'НАЗВ_ОТД'],
    rows=[
        [1, 'IT-отдел'],
        [2, 'Аналитика'],
    ],
    key_cols={0},
    caption='Таблица 6. Отношение ОТДЕЛЫ'
)

doc.add_paragraph()
p = doc.add_paragraph()
r1 = p.add_run('3. ВАКАНСИИ')
set_font(r1, 12, bold=True)
r2 = p.add_run('  (Первичный ключ: ')
set_font(r2, 12)
r3 = p.add_run('Н_ВАК')
set_font(r3, 12, bold=True, color=(255, 0, 0))
r4 = p.add_run(')')
set_font(r4, 12)
add_table(
    headers=['Н_ВАК', 'НАЗВ_ВАК', 'Н_ОТД'],
    rows=[
        [1, 'Разработчик', 1],
        [2, 'Аналитик', 2],
    ],
    key_cols={0},
    caption='Таблица 7. Отношение ВАКАНСИИ'
)

doc.add_paragraph()
p = doc.add_paragraph()
r1 = p.add_run('4. МЕНЕДЖЕРЫ')
set_font(r1, 12, bold=True)
r2 = p.add_run('  (Первичный ключ: ')
set_font(r2, 12)
r3 = p.add_run('Н_МЕН')
set_font(r3, 12, bold=True, color=(255, 0, 0))
r4 = p.add_run(')')
set_font(r4, 12)
add_table(
    headers=['Н_МЕН', 'ФИО_МЕН'],
    rows=[
        [1, 'Сидоров А.А.'],
        [2, 'Петрова В.В.'],
    ],
    key_cols={0},
    caption='Таблица 8. Отношение МЕНЕДЖЕРЫ'
)

doc.add_paragraph()
p = doc.add_paragraph()
r1 = p.add_run('5. СОБЕСЕДОВАНИЯ')
set_font(r1, 12, bold=True)
r2 = p.add_run('  (Первичный ключ: ')
set_font(r2, 12)
r3 = p.add_run('{Н_КАНД, Н_ВАК}')
set_font(r3, 12, bold=True, color=(255, 0, 0))
r4 = p.add_run(')')
set_font(r4, 12)
add_table(
    headers=['Н_КАНД', 'Н_ВАК', 'Н_МЕН', 'ДАТА', 'РЕЗУЛЬТАТ'],
    rows=[
        [1, 1, 1, '10.10.23', 'Оффер'],
        [1, 2, 2, '12.10.23', 'Отказ'],
        [2, 1, 1, '13.10.23', 'В резерв'],
    ],
    key_cols={0, 1},
    caption='Таблица 9. Отношение СОБЕСЕДОВАНИЯ'
)

# ============================================================
# ИТОГОВАЯ БД
# ============================================================
doc.add_paragraph()
add_heading('Итоговая БД', level=1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Cm(1.25)
r = p.add_run(
    'Исходное отношение было нормализовано — приведено к третьей нормальной форме за счёт '
    'декомпозиции в пять взаимосвязанных справочников:'
)
set_font(r, 12)

final = [
    'КАНДИДАТЫ (Ключ: Н_КАНД)',
    'ОТДЕЛЫ (Ключ: Н_ОТД)',
    'ВАКАНСИИ (Ключ: Н_ВАК)',
    'МЕНЕДЖЕРЫ (Ключ: Н_МЕН)',
    'СОБЕСЕДОВАНИЯ (Ключ: {Н_КАНД, Н_ВАК})',
]
for item in final:
    add_bullet(item)

doc.save('/home/user/python_hse_TheFirstCours/HW1.docx')
print('Done: HW1.docx created')
